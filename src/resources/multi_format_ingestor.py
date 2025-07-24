"""
Multi-format resource ingestion system.
Handles PDFs, DOCX, URLs, text files, and other document formats.
"""

import streamlit as st
import json
import requests
from pathlib import Path
from typing import Dict, List, Any, Optional, Union
from dataclasses import dataclass, asdict
import hashlib
from datetime import datetime
import re
import tempfile
import os

# Document processing imports
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False


@dataclass
class ResourceMetadata:
    """Metadata for a resource."""
    id: str
    title: str
    source_type: str  # 'pdf', 'docx', 'url', 'text', 'video'
    source_url: str
    author: Optional[str] = None
    date_added: str = None
    file_size: Optional[int] = None
    page_count: Optional[int] = None
    description: Optional[str] = None
    tags: List[str] = None
    
    def __post_init__(self):
        if self.date_added is None:
            self.date_added = datetime.now().isoformat()
        if self.tags is None:
            self.tags = []


@dataclass
class ProcessedResource:
    """A processed resource with extracted content."""
    metadata: ResourceMetadata
    content: str
    chunks: List[str]
    chunk_count: int
    processing_notes: List[str] = None
    
    def __post_init__(self):
        if self.processing_notes is None:
            self.processing_notes = []


class MultiFormatIngestor:
    """Ingest and process multiple document formats."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.supported_formats = self._get_supported_formats()
    
    def _get_supported_formats(self) -> Dict[str, bool]:
        """Get list of supported formats based on available libraries."""
        return {
            'pdf': PDF_AVAILABLE,
            'docx': DOCX_AVAILABLE,
            'url': BS4_AVAILABLE,
            'text': True,
            'markdown': MARKDOWN_AVAILABLE,
            'json': True,
            'csv': True
        }
    
    def get_missing_dependencies(self) -> List[str]:
        """Get list of missing dependencies for full functionality."""
        missing = []
        if not PDF_AVAILABLE:
            missing.append("PyPDF2 (for PDF processing)")
        if not DOCX_AVAILABLE:
            missing.append("python-docx (for DOCX processing)")
        if not BS4_AVAILABLE:
            missing.append("beautifulsoup4 (for web scraping)")
        if not MARKDOWN_AVAILABLE:
            missing.append("markdown (for Markdown processing)")
        return missing
    
    def process_pdf(self, file_path: str, metadata: ResourceMetadata) -> ProcessedResource:
        """Process PDF file."""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not available. Install with: pip install PyPDF2")
        
        content = ""
        processing_notes = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                page_count = len(pdf_reader.pages)
                metadata.page_count = page_count
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        content += f"\n--- Page {page_num} ---\n{page_text}\n"
                    except Exception as e:
                        processing_notes.append(f"Error extracting page {page_num}: {str(e)}")
                        
        except Exception as e:
            processing_notes.append(f"Error reading PDF: {str(e)}")
            content = f"Error processing PDF: {str(e)}"
        
        chunks = self._chunk_text(content)
        
        return ProcessedResource(
            metadata=metadata,
            content=content,
            chunks=chunks,
            chunk_count=len(chunks),
            processing_notes=processing_notes
        )
    
    def process_docx(self, file_path: str, metadata: ResourceMetadata) -> ProcessedResource:
        """Process DOCX file."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available. Install with: pip install python-docx")
        
        content = ""
        processing_notes = []
        
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    content += para.text + "\n\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    content += row_text + "\n"
                content += "\n"
            
            processing_notes.append(f"Extracted {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
            
        except Exception as e:
            processing_notes.append(f"Error reading DOCX: {str(e)}")
            content = f"Error processing DOCX: {str(e)}"
        
        chunks = self._chunk_text(content)
        
        return ProcessedResource(
            metadata=metadata,
            content=content,
            chunks=chunks,
            chunk_count=len(chunks),
            processing_notes=processing_notes
        )
    
    def process_url(self, url: str, metadata: ResourceMetadata) -> ProcessedResource:
        """Process web URL."""
        if not BS4_AVAILABLE:
            raise ImportError("beautifulsoup4 not available. Install with: pip install beautifulsoup4")
        
        content = ""
        processing_notes = []
        
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            
            response = requests.get(url, headers=headers, timeout=30)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Try to extract title
            title_tag = soup.find('title')
            if title_tag and not metadata.title:
                metadata.title = title_tag.get_text().strip()
            
            # Extract main content
            # Try common content containers first
            main_content = None
            for selector in ['main', 'article', '.content', '#content', '.post', '.entry']:
                main_content = soup.select_one(selector)
                if main_content:
                    break
            
            if main_content:
                content = main_content.get_text(separator='\n', strip=True)
                processing_notes.append("Extracted from main content area")
            else:
                # Fall back to body
                body = soup.find('body')
                if body:
                    content = body.get_text(separator='\n', strip=True)
                    processing_notes.append("Extracted from full body")
                else:
                    content = soup.get_text(separator='\n', strip=True)
                    processing_notes.append("Extracted from entire page")
            
            # Clean up the text
            content = re.sub(r'\n\s*\n', '\n\n', content)  # Remove excessive newlines
            content = re.sub(r'[ \t]+', ' ', content)  # Normalize spaces
            
            processing_notes.append(f"Content length: {len(content)} characters")
            
        except requests.RequestException as e:
            processing_notes.append(f"Error fetching URL: {str(e)}")
            content = f"Error processing URL: {str(e)}"
        except Exception as e:
            processing_notes.append(f"Error parsing content: {str(e)}")
            content = f"Error processing URL content: {str(e)}"
        
        chunks = self._chunk_text(content)
        
        return ProcessedResource(
            metadata=metadata,
            content=content,
            chunks=chunks,
            chunk_count=len(chunks),
            processing_notes=processing_notes
        )
    
    def process_text_file(self, file_path: str, metadata: ResourceMetadata) -> ProcessedResource:
        """Process plain text file."""
        content = ""
        processing_notes = []
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                    processing_notes.append(f"Successfully read with {encoding} encoding")
                    break
                except UnicodeDecodeError:
                    continue
            
            if not content:
                processing_notes.append("Could not decode file with any standard encoding")
                content = "Error: Could not decode file content"
            
        except Exception as e:
            processing_notes.append(f"Error reading text file: {str(e)}")
            content = f"Error processing text file: {str(e)}"
        
        chunks = self._chunk_text(content)
        
        return ProcessedResource(
            metadata=metadata,
            content=content,
            chunks=chunks,
            chunk_count=len(chunks),
            processing_notes=processing_notes
        )
    
    def process_markdown(self, file_path: str, metadata: ResourceMetadata) -> ProcessedResource:
        """Process Markdown file."""
        processing_notes = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
            
            if MARKDOWN_AVAILABLE:
                # Convert to HTML then extract text
                html = markdown.markdown(md_content)
                soup = BeautifulSoup(html, 'html.parser')
                content = soup.get_text(separator='\n', strip=True)
                processing_notes.append("Converted from Markdown to text")
            else:
                # Use raw markdown
                content = md_content
                processing_notes.append("Processed as raw markdown (markdown library not available)")
                
        except Exception as e:
            processing_notes.append(f"Error reading markdown file: {str(e)}")
            content = f"Error processing markdown file: {str(e)}"
        
        chunks = self._chunk_text(content)
        
        return ProcessedResource(
            metadata=metadata,
            content=content,
            chunks=chunks,
            chunk_count=len(chunks),
            processing_notes=processing_notes
        )
    
    def _chunk_text(self, text: str) -> List[str]:
        """Split text into chunks."""
        if not text or len(text) <= self.chunk_size:
            return [text] if text else []
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            if end >= len(text):
                # Last chunk
                chunks.append(text[start:])
                break
            
            # Try to break at sentence boundary
            chunk_text = text[start:end]
            
            # Look for sentence endings in the last part of the chunk
            last_part = chunk_text[-200:] if len(chunk_text) > 200 else chunk_text
            sentence_ends = [m.end() for m in re.finditer(r'[.!?]\s+', last_part)]
            
            if sentence_ends:
                # Break at the last sentence in the chunk
                break_point = start + len(chunk_text) - len(last_part) + sentence_ends[-1]
                chunks.append(text[start:break_point])
                start = break_point - self.chunk_overlap
            else:
                # No good break point, break at word boundary
                space_index = chunk_text.rfind(' ')
                if space_index > self.chunk_size * 0.8:  # Don't break too early
                    chunks.append(text[start:start + space_index])
                    start = start + space_index - self.chunk_overlap
                else:
                    # Just break at character boundary
                    chunks.append(chunk_text)
                    start = end - self.chunk_overlap
            
            # Ensure we're making progress
            if start < 0:
                start = 0
        
        return [chunk.strip() for chunk in chunks if chunk.strip()]
    
    def process_resource(self, source: Union[str, bytes], source_type: str, metadata: ResourceMetadata) -> ProcessedResource:
        """Process a resource based on its type."""
        
        if source_type == 'pdf':
            if isinstance(source, str):  # File path
                return self.process_pdf(source, metadata)
            else:  # Bytes content
                with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_file:
                    tmp_file.write(source)
                    tmp_file.flush()
                    result = self.process_pdf(tmp_file.name, metadata)
                    os.unlink(tmp_file.name)
                    return result
        
        elif source_type == 'docx':
            if isinstance(source, str):  # File path
                return self.process_docx(source, metadata)
            else:  # Bytes content
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                    tmp_file.write(source)
                    tmp_file.flush()
                    result = self.process_docx(tmp_file.name, metadata)
                    os.unlink(tmp_file.name)
                    return result
        
        elif source_type == 'url':
            return self.process_url(source, metadata)
        
        elif source_type == 'text':
            if isinstance(source, str):
                return self.process_text_file(source, metadata)
            else:
                # Direct text content
                content = source.decode('utf-8') if isinstance(source, bytes) else source
                chunks = self._chunk_text(content)
                return ProcessedResource(
                    metadata=metadata,
                    content=content,
                    chunks=chunks,
                    chunk_count=len(chunks)
                )
        
        elif source_type == 'markdown':
            return self.process_markdown(source, metadata)
        
        else:
            raise ValueError(f"Unsupported source type: {source_type}")


class ResourceManager:
    """Manage the unified knowledge base with multiple resource types."""
    
    def __init__(self, data_directory: str = "data"):
        self.data_directory = Path(data_directory)
        self.data_directory.mkdir(exist_ok=True)
        
        self.knowledge_file = self.data_directory / "unified_knowledge_base.json"
        self.ingestor = MultiFormatIngestor()
        
        # Load existing knowledge base
        self.knowledge_base = self._load_knowledge_base()
    
    def _load_knowledge_base(self) -> Dict[str, Any]:
        """Load the unified knowledge base."""
        if self.knowledge_file.exists():
            try:
                with open(self.knowledge_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception as e:
                st.error(f"Error loading knowledge base: {e}")
                return {}
        
        # If no unified knowledge base exists, try to load the original video knowledge base
        original_kb_file = Path("knowledge_base_final.json")
        if original_kb_file.exists():
            try:
                with open(original_kb_file, 'r', encoding='utf-8') as f:
                    original_kb = json.load(f)
                
                # Convert video knowledge base to new format
                unified_kb = {}
                for url, video_data in original_kb.items():
                    resource_id = self._generate_id(url)
                    unified_kb[resource_id] = {
                        'metadata': {
                            'id': resource_id,
                            'title': video_data['title'],
                            'source_type': 'video',
                            'source_url': url,
                            'author': video_data.get('uploader', 'Unknown'),
                            'date_added': datetime.now().isoformat(),
                            'description': video_data.get('description', ''),
                            'tags': ['video', 'ai', 'education']
                        },
                        'content': video_data.get('transcript', ''),
                        'chunks': video_data.get('chunks', []),
                        'chunk_count': len(video_data.get('chunks', [])),
                        'processing_notes': ['Imported from original video knowledge base']
                    }
                
                # Save the converted knowledge base
                self._save_knowledge_base(unified_kb)
                return unified_kb
                
            except Exception as e:
                st.error(f"Error converting original knowledge base: {e}")
                return {}
        
        return {}
    
    def _save_knowledge_base(self, kb: Dict[str, Any] = None):
        """Save the knowledge base to file."""
        if kb is None:
            kb = self.knowledge_base
        
        try:
            with open(self.knowledge_file, 'w', encoding='utf-8') as f:
                json.dump(kb, f, indent=2, ensure_ascii=False)
        except Exception as e:
            st.error(f"Error saving knowledge base: {e}")
    
    def _generate_id(self, source: str) -> str:
        """Generate a unique ID for a resource."""
        return hashlib.md5(source.encode()).hexdigest()[:12]
    
    def add_resource(self, source: Union[str, bytes], source_type: str, title: str = None, 
                    author: str = None, description: str = None, tags: List[str] = None) -> bool:
        """Add a new resource to the knowledge base."""
        
        try:
            # Create metadata
            resource_id = self._generate_id(str(source))
            
            metadata = ResourceMetadata(
                id=resource_id,
                title=title or f"Resource {resource_id}",
                source_type=source_type,
                source_url=str(source),
                author=author,
                description=description,
                tags=tags or []
            )
            
            # Process the resource
            processed_resource = self.ingestor.process_resource(source, source_type, metadata)
            
            # Add to knowledge base
            self.knowledge_base[resource_id] = {
                'metadata': asdict(processed_resource.metadata),
                'content': processed_resource.content,
                'chunks': processed_resource.chunks,
                'chunk_count': processed_resource.chunk_count,
                'processing_notes': processed_resource.processing_notes
            }
            
            # Save to file
            self._save_knowledge_base()
            
            return True
            
        except Exception as e:
            st.error(f"Error adding resource: {e}")
            return False
    
    def remove_resource(self, resource_id: str) -> bool:
        """Remove a resource from the knowledge base."""
        if resource_id in self.knowledge_base:
            del self.knowledge_base[resource_id]
            self._save_knowledge_base()
            return True
        return False
    
    def search_resources(self, query: str, n_results: int = 5, source_types: List[str] = None) -> List[Dict[str, Any]]:
        """Search across all resources."""
        query_words = query.lower().split()
        results = []
        
        for resource_id, resource_data in self.knowledge_base.items():
            metadata = resource_data['metadata']
            
            # Filter by source type if specified
            if source_types and metadata['source_type'] not in source_types:
                continue
            
            title = metadata['title'].lower()
            content = resource_data['content'].lower()
            
            # Score based on keyword matches
            score = 0
            for word in query_words:
                score += title.count(word) * 3
                score += content.count(word)
                score += sum(tag.lower().count(word) for tag in metadata.get('tags', []))
            
            if score > 0:
                # Find best matching chunk
                chunks = resource_data.get('chunks', [])
                best_chunk = ""
                
                if chunks:
                    best_chunk = chunks[0]
                    for chunk in chunks[:3]:
                        chunk_lower = chunk.lower()
                        if any(word in chunk_lower for word in query_words):
                            best_chunk = chunk
                            break
                else:
                    best_chunk = content[:500]
                
                results.append({
                    'content': best_chunk,
                    'metadata': metadata,
                    'distance': 1.0 - (score / 100),  # Convert score to distance
                    'resource_id': resource_id
                })
        
        # Sort by score (lower distance = better match)
        results.sort(key=lambda x: x['distance'])
        return results[:n_results]
    
    def get_resource_stats(self) -> Dict[str, Any]:
        """Get statistics about the knowledge base."""
        total_resources = len(self.knowledge_base)
        
        # Count by type and separate Daily.dev articles
        type_counts = {}
        daily_dev_count = 0
        total_chunks = 0
        
        for resource_data in self.knowledge_base.values():
            source_type = resource_data['metadata']['source_type']
            tags = resource_data['metadata'].get('tags', [])
            source_url = resource_data['metadata'].get('source_url', '')
            
            # Check if this is a Daily.dev/tech article
            is_daily_dev = 'daily.dev' in tags or 'tech' in tags or 'daily.dev' in source_url
            
            if is_daily_dev:
                # Count as daily.dev instead of url
                type_counts['daily.dev'] = type_counts.get('daily.dev', 0) + 1
                daily_dev_count += 1
            else:
                # Count normally
                type_counts[source_type] = type_counts.get(source_type, 0) + 1
            
            total_chunks += resource_data.get('chunk_count', 0)
        
        # Enhanced type counts is the same as type_counts now
        enhanced_type_counts = type_counts.copy()
        
        return {
            'total_resources': total_resources,
            'total_chunks': total_chunks,
            'by_type': type_counts,
            'enhanced_by_type': enhanced_type_counts,
            'daily_dev_count': daily_dev_count,
            'supported_formats': self.ingestor.supported_formats,
            'missing_dependencies': self.ingestor.get_missing_dependencies()
        }